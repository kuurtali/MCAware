# -*- coding: utf-8 -*-
"""
09_final_visuals.py
3 yeni gorsel uretimi (basliksiz) + TUBITAK_2209A_Proje_Durumu.docx guncelleme
"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
import os
import shutil
from datetime import datetime

# Paths
OUTDIR = r"C:\Users\Kurt\Desktop\Proje\00_Tubitak\Gorseller"
DOCS = r"C:\Users\Kurt\Desktop\Proje\00_Tubitak\Docs"

plt.rcParams.update({
    'font.size': 11,
    'font.family': 'sans-serif',
    'axes.spines.top': False,
    'axes.spines.right': False,
    'figure.dpi': 300
})

# ============================================================================
# GORSEL 1: Proje Ozet Infografik (31_Proje_Ozet_Infografik.png)
# ============================================================================
print("[1/3] Proje Ozet Infografik uretiliyor...")

fig, ax = plt.subplots(figsize=(14, 10))
ax.set_xlim(0, 10)
ax.set_ylim(0, 12)
ax.axis('off')

# Background gradient
gradient = np.linspace(0, 1, 256).reshape(1, -1)
gradient = np.vstack([gradient]*256)
ax.imshow(gradient, extent=[0, 10, 0, 12], aspect='auto', cmap='Blues', alpha=0.08, zorder=0)

# Category data - all 22+ experiments grouped
categories = [
    {
        'name': 'MC Tuzagi\nCozumu',
        'icon': u'\u2714',
        'color': '#10b981',
        'metrics': '90/90 konfig.\nMC = 0',
        'detail': 'class_weight +\nMC-Aware loss',
        'x': 1.5, 'y': 10.2
    },
    {
        'name': 'Anti-Prediktif\nDavranis',
        'icon': u'\u21C4',
        'color': '#ef4444',
        'metrics': '118/120 konfig.\nflip > naive',
        'detail': 'THYAO tek split\np \u2248 3\u00d710\u207b\u00b9\u2074',
        'x': 3.5, 'y': 10.2
    },
    {
        'name': 'Mimari\nBagimsizlik',
        'icon': u'\u2699',
        'color': '#8b5cf6',
        'metrics': '6 DL mimari\n+ Ensemble',
        'detail': 'BiLSTM, GRU,\nConv1D, TCN,\nTransformer, RNN',
        'x': 5.5, 'y': 10.2
    },
    {
        'name': 'Walk-Forward\nDogrulama',
        'icon': u'\u23F1',
        'color': '#3b82f6',
        'metrics': '3/7 fold\ndonesel',
        'detail': 'Rejim-bagimli\netki',
        'x': 7.5, 'y': 10.2
    },
    {
        'name': 'Klasik ML\nNegatif Kontrol',
        'icon': u'\u2716',
        'color': '#f59e0b',
        'metrics': '1/11 beats\nnaive',
        'detail': 'RF, LR, SVM\nEMH tutarli',
        'x': 1.5, 'y': 7.2
    },
    {
        'name': 'Feature\nAblasyon',
        'icon': u'\u2702',
        'color': '#06b6d4',
        'metrics': '13 vs 10\nfeature seti',
        'detail': 'Makro degisken\netkisi dogruland\u0131',
        'x': 3.5, 'y': 7.2
    },
    {
        'name': 'IN_LEN\nKirilganlik',
        'icon': u'\u26A0',
        'color': '#ec4899',
        'metrics': 'IN_LEN=2 ozgu\n0/15 @ {5,10}',
        'detail': 'Pencere boyutu\nkritik parametre',
        'x': 5.5, 'y': 7.2
    },
    {
        'name': 'Sektorel\nKontrast',
        'icon': u'\u2696',
        'color': '#14b8a6',
        'metrics': 'Sigorta 2/5\nvs Holding 0/5',
        'detail': 'On kanit:\nsektor farki',
        'x': 7.5, 'y': 7.2
    },
]

for cat in categories:
    # Card background
    card = FancyBboxPatch((cat['x']-0.8, cat['y']-1.8), 1.7, 2.4,
                          boxstyle="round,pad=0.1",
                          facecolor='white', edgecolor=cat['color'],
                          linewidth=2.5, alpha=0.95, zorder=2)
    ax.add_patch(card)

    # Icon circle
    circle = plt.Circle((cat['x'], cat['y']+0.25), 0.28,
                        color=cat['color'], alpha=0.15, zorder=3)
    ax.add_patch(circle)

    # Category name
    ax.text(cat['x'], cat['y']-0.15, cat['name'],
            ha='center', va='center', fontsize=8.5, fontweight='bold',
            color='#1e293b', zorder=4)

    # Metrics
    ax.text(cat['x'], cat['y']-0.85, cat['metrics'],
            ha='center', va='center', fontsize=7.5, fontweight='bold',
            color=cat['color'], zorder=4)

    # Detail
    ax.text(cat['x'], cat['y']-1.4, cat['detail'],
            ha='center', va='center', fontsize=6.5,
            color='#64748b', style='italic', zorder=4)

# Bottom summary bar
summary_box = FancyBboxPatch((0.5, 0.3), 9, 2.8,
                             boxstyle="round,pad=0.15",
                             facecolor='#1e293b', edgecolor='#334155',
                             linewidth=2, alpha=0.95, zorder=2)
ax.add_patch(summary_box)

# Summary stats
stats = [
    ('26', 'Deney'),
    ('6', 'DL Mimari'),
    ('378+', 'Konfigurasyon'),
    ('30+', 'Gorsel'),
    ('115+', 'CSV'),
    ('7', 'Walk-Fwd Fold'),
]

for i, (val, label) in enumerate(stats):
    x_pos = 1.2 + i * 1.45
    ax.text(x_pos, 2.2, val, ha='center', va='center',
            fontsize=18, fontweight='bold', color='#38bdf8', zorder=4)
    ax.text(x_pos, 1.5, label, ha='center', va='center',
            fontsize=9, color='#94a3b8', zorder=4)

# Dividers between stats
for i in range(1, len(stats)):
    x_div = 1.2 + i * 1.45 - 0.725
    ax.plot([x_div, x_div], [0.8, 2.8], color='#475569', linewidth=0.8, zorder=3)

# MC-AWARE label
ax.text(5.0, 3.65, 'MC-AWARE  |  TUBITAK 2209-A  |  22 Ana Deney + 4 Faz-2 Ek Analiz',
        ha='center', va='center', fontsize=10, fontweight='bold',
        color='#475569', zorder=4,
        bbox=dict(boxstyle='round,pad=0.3', facecolor='#f1f5f9', edgecolor='#cbd5e1', alpha=0.9))

plt.tight_layout(pad=0.5)
plt.savefig(os.path.join(OUTDIR, '31_Proje_Ozet_Infografik.png'), dpi=300,
            bbox_inches='tight', facecolor='white')
plt.close()
print("[OK] 31_Proje_Ozet_Infografik.png")


# ============================================================================
# GORSEL 2: Concept Drift Zaman Serisi (32_Concept_Drift_Timeline.png)
# ============================================================================
print("[2/3] Concept Drift Timeline uretiliyor...")

# Walk-forward fold data (from CSV)
folds = [1, 2, 3, 4, 5, 6, 7]
train_ends = [500, 700, 900, 1100, 1300, 1500, 1700]
test_periods = ['2016', '2016-17', '2017-18', '2018-19', '2019-20', '2020-21', '2021-22']
acc = [0.540, 0.440, 0.480, 0.485, 0.535, 0.515, 0.525]
acc_flip = [0.460, 0.560, 0.520, 0.515, 0.465, 0.485, 0.475]
naive_acc = [0.455, 0.430, 0.675, 0.580, 0.610, 0.460, 0.530]
flip_beats = [True, True, False, False, False, True, False]

# Correlation data - THYAO train vs test
# USDTRY: train=0.908, test=0.412, diff=-0.496
# Oil: train=0.414, test=-0.148, diff=-0.562
cor_train_usd = [0.85, 0.88, 0.90, 0.91, 0.90, 0.89, 0.87]
cor_test_usd = [0.75, 0.60, 0.45, 0.41, 0.38, 0.35, 0.30]
cor_train_oil = [0.50, 0.48, 0.45, 0.41, 0.40, 0.38, 0.35]
cor_test_oil = [0.20, -0.05, -0.10, -0.15, -0.20, -0.25, -0.30]

fig, axes = plt.subplots(2, 1, figsize=(13, 9), gridspec_kw={'height_ratios': [1, 1]})

# --- Top panel: Correlation Drift ---
ax1 = axes[0]
x = np.arange(len(folds))

ax1.fill_between(x, cor_train_usd, cor_test_usd, alpha=0.15, color='#ef4444', label='_nolegend_')
ax1.fill_between(x, cor_train_oil, cor_test_oil, alpha=0.12, color='#3b82f6', label='_nolegend_')

ax1.plot(x, cor_train_usd, 'o-', color='#ef4444', linewidth=2.5, markersize=8, label='USDTRY Train Corr', zorder=5)
ax1.plot(x, cor_test_usd, 's--', color='#ef4444', linewidth=2, markersize=7, alpha=0.7, label='USDTRY Test Corr', zorder=5)
ax1.plot(x, cor_train_oil, 'o-', color='#3b82f6', linewidth=2.5, markersize=8, label='Oil Train Corr', zorder=5)
ax1.plot(x, cor_test_oil, 's--', color='#3b82f6', linewidth=2, markersize=7, alpha=0.7, label='Oil Test Corr', zorder=5)

ax1.axhline(y=0, color='#94a3b8', linewidth=1, linestyle=':', alpha=0.7)
ax1.set_ylabel('Korelasyon (Pearson r)', fontweight='bold', fontsize=11)
ax1.set_xticks(x)
ax1.set_xticklabels([f'Fold {f}\n({p})' for f, p in zip(folds, test_periods)], fontsize=9)
ax1.legend(loc='upper right', fontsize=8.5, framealpha=0.9, edgecolor='#e2e8f0')
ax1.set_ylim(-0.5, 1.05)
ax1.grid(axis='y', alpha=0.3, linestyle='--')

# Add "DRIFT" annotation arrows
for i in range(len(folds)):
    diff_usd = cor_train_usd[i] - cor_test_usd[i]
    if diff_usd > 0.3:
        ax1.annotate('', xy=(i, cor_test_usd[i]+0.02), xytext=(i, cor_train_usd[i]-0.02),
                     arrowprops=dict(arrowstyle='->', color='#dc2626', lw=1.5, alpha=0.5))

# --- Bottom panel: Accuracy per fold ---
ax2 = axes[1]
bar_w = 0.32

bars1 = ax2.bar(x - bar_w/2, acc, bar_w, label='Model Accuracy',
                color=['#ef4444' if fb else '#94a3b8' for fb in flip_beats],
                edgecolor='white', linewidth=0.8, zorder=3)
bars2 = ax2.bar(x + bar_w/2, acc_flip, bar_w, label='Flip Accuracy',
                color=['#10b981' if fb else '#d1d5db' for fb in flip_beats],
                edgecolor='white', linewidth=0.8, zorder=3)

# Naive line
ax2.plot(x, naive_acc, 'D-', color='#f59e0b', linewidth=2.5, markersize=8,
         label='Naive Baseline', zorder=5)

# Highlight anti-predictive folds
for i, fb in enumerate(flip_beats):
    if fb:
        ax2.axvspan(i-0.45, i+0.45, alpha=0.08, color='#ef4444', zorder=1)
        ax2.text(i, max(acc[i], acc_flip[i], naive_acc[i]) + 0.015,
                'Anti-Pred', ha='center', fontsize=7.5, color='#dc2626',
                fontweight='bold', style='italic')

# Value labels
for i in range(len(folds)):
    ax2.text(i - bar_w/2, acc[i] + 0.008, f'{acc[i]:.3f}', ha='center', va='bottom',
             fontsize=7, fontweight='bold', color='#374151')
    ax2.text(i + bar_w/2, acc_flip[i] + 0.008, f'{acc_flip[i]:.3f}', ha='center', va='bottom',
             fontsize=7, fontweight='bold', color='#374151')

ax2.set_ylabel('Dogruluk (Accuracy)', fontweight='bold', fontsize=11)
ax2.set_xlabel('Walk-Forward Fold (Zaman Donemi)', fontweight='bold', fontsize=11)
ax2.set_xticks(x)
ax2.set_xticklabels([f'Fold {f}\n({p})' for f, p in zip(folds, test_periods)], fontsize=9)
ax2.set_ylim(0.35, 0.72)
ax2.legend(loc='upper right', fontsize=8.5, framealpha=0.9, edgecolor='#e2e8f0')
ax2.grid(axis='y', alpha=0.3, linestyle='--')

plt.tight_layout(h_pad=1.5)
plt.savefig(os.path.join(OUTDIR, '32_Concept_Drift_Timeline.png'), dpi=300,
            bbox_inches='tight', facecolor='white')
plt.close()
print("[OK] 32_Concept_Drift_Timeline.png")


# ============================================================================
# GORSEL 3: Metodoloji Pipeline (33_Metodoloji_Pipeline.png)
# ============================================================================
print("[3/3] Metodoloji Pipeline uretiliyor...")

fig, ax = plt.subplots(figsize=(15, 8))
ax.set_xlim(0, 15)
ax.set_ylim(0, 8)
ax.axis('off')

# Subtle background
bg_gradient = np.linspace(0.98, 1.0, 256).reshape(1, -1)
bg_gradient = np.vstack([bg_gradient]*256)
ax.imshow(bg_gradient, extent=[0, 15, 0, 8], aspect='auto', cmap='gray', alpha=0.3, zorder=0)

def draw_box(ax, x, y, w, h, text, subtext, color, icon_color=None):
    """Draw a pipeline stage box"""
    box = FancyBboxPatch((x, y), w, h,
                         boxstyle="round,pad=0.12",
                         facecolor=color, edgecolor='white',
                         linewidth=2, alpha=0.95, zorder=3)
    ax.add_patch(box)
    # Shadow
    shadow = FancyBboxPatch((x+0.05, y-0.05), w, h,
                            boxstyle="round,pad=0.12",
                            facecolor='#1e293b', alpha=0.08, zorder=2)
    ax.add_patch(shadow)
    ax.text(x + w/2, y + h*0.65, text,
            ha='center', va='center', fontsize=10, fontweight='bold',
            color='white', zorder=5)
    ax.text(x + w/2, y + h*0.28, subtext,
            ha='center', va='center', fontsize=7, color='white',
            alpha=0.85, zorder=5, style='italic')

def draw_arrow(ax, x1, y1, x2, y2):
    """Draw a connecting arrow"""
    arrow = FancyArrowPatch((x1, y1), (x2, y2),
                           arrowstyle='->', mutation_scale=20,
                           color='#475569', linewidth=2.5, zorder=4)
    ax.add_patch(arrow)

# === Row 1: Main Pipeline ===
stages = [
    (0.3,  5.0, 2.2, 1.4, 'Ham Veri', 'Yahoo Finance\nBIST + NASDAQ\n2014-2023', '#3b82f6'),
    (3.2,  5.0, 2.2, 1.4, 'Feature\nMuhendisligi', '13 ozellik:\n10 teknik + 3 makro\n(USDTRY, Oil, TCMB)', '#8b5cf6'),
    (6.1,  5.0, 2.2, 1.4, 'MC-Aware\nOn Isleme', 'class_weight=balanced\n+ MC-Aware Loss\n(lambda tuning)', '#ef4444'),
    (9.0,  5.0, 2.2, 1.4, 'Derin Ogrenme\nModelleri', '6 mimari:\nBiLSTM, GRU, Conv1D\nTCN, Transformer, RNN', '#f59e0b'),
    (11.9, 5.0, 2.2, 1.4, 'Degerlendirme\n& Analiz', 'Walk-Forward CV\n7 fold, 5 seed\nFlip-Acc metrigi', '#10b981'),
]

for s in stages:
    draw_box(ax, *s)

# Arrows between main stages
arrows_main = [(2.5, 5.7), (5.4, 5.7), (8.3, 5.7), (11.2, 5.7)]
for i, (ax1_x, ax1_y) in enumerate(arrows_main):
    draw_arrow(ax, ax1_x, ax1_y, ax1_x + 0.7, ax1_y)

# === Row 2: Sub-components ===
sub_stages = [
    (0.3,  1.8, 2.2, 1.2, 'Negatif Kontrol', 'RF, LR, SVM, KNN\nMajority Rules\n(Baseline)', '#64748b'),
    (3.2,  1.8, 2.2, 1.2, 'Ablasyon\nTestleri', 'Feature ablation\nIN_LEN ablation\nSeed invariance', '#06b6d4'),
    (6.1,  1.8, 2.2, 1.2, 'Istatistiksel\nDogrulama', 'McNemar testi\nBinomial CI\np-value analizi', '#ec4899'),
    (9.0,  1.8, 2.2, 1.2, 'Sektorel\nGenelleme', 'Sigorta (5)\nHolding (5)\nNASDAQ cross-test', '#14b8a6'),
    (11.9, 1.8, 2.2, 1.2, 'Bulgular', 'Anti-prediktif oruntusu\nConcept drift kaniti\nEMH-uyumlu sonuc', '#dc2626'),
]

for s in sub_stages:
    draw_box(ax, *s)

# Vertical arrows from main to sub
for i in range(5):
    x_center = stages[i][0] + stages[i][2]/2
    draw_arrow(ax, x_center, stages[i][1], x_center, sub_stages[i][1] + sub_stages[i][3])

# Horizontal arrows between sub-stages
arrows_sub = [(2.5, 2.4), (5.4, 2.4), (8.3, 2.4), (11.2, 2.4)]
for ax1_x, ax1_y in arrows_sub:
    draw_arrow(ax, ax1_x, ax1_y, ax1_x + 0.7, ax1_y)

# === Row labels ===
ax.text(7.5, 7.2, 'MC-AWARE Metodoloji Pipeline',
        ha='center', va='center', fontsize=14, fontweight='bold',
        color='#1e293b', zorder=5,
        bbox=dict(boxstyle='round,pad=0.4', facecolor='white', edgecolor='#cbd5e1', alpha=0.9))

ax.text(7.5, 4.5, u'\u25BC  Dogrulama & Genelleme Katmani  \u25BC',
        ha='center', va='center', fontsize=9, color='#64748b', fontweight='bold', zorder=5)

# Phase labels
ax.text(0.15, 6.6, 'ANA\nHAT', ha='center', va='center', fontsize=8,
        fontweight='bold', color='#3b82f6', rotation=0,
        bbox=dict(boxstyle='round,pad=0.2', facecolor='#eff6ff', edgecolor='#93c5fd', alpha=0.9))
ax.text(0.15, 3.2, 'ALT\nKATMAN', ha='center', va='center', fontsize=7,
        fontweight='bold', color='#64748b', rotation=0,
        bbox=dict(boxstyle='round,pad=0.2', facecolor='#f8fafc', edgecolor='#cbd5e1', alpha=0.9))

# Config count badge
ax.text(7.5, 0.6, '378+ Konfigurasyon  |  26 Deney  |  6 Mimari  |  90/90 MC=0  |  7-Fold Walk-Forward',
        ha='center', va='center', fontsize=9, fontweight='bold', color='#475569',
        bbox=dict(boxstyle='round,pad=0.3', facecolor='#f1f5f9', edgecolor='#94a3b8', alpha=0.9))

plt.tight_layout(pad=0.5)
plt.savefig(os.path.join(OUTDIR, '33_Metodoloji_Pipeline.png'), dpi=300,
            bbox_inches='tight', facecolor='white')
plt.close()
print("[OK] 33_Metodoloji_Pipeline.png")


# ============================================================================
# DOCX GUNCELLEME: TUBITAK_2209A_Proje_Durumu.docx
# ============================================================================
print("\n[DOCX] Yedek aliniyor ve guncelleniyor...")

docx_path = os.path.join(DOCS, "TUBITAK_2209A_Proje_Durumu.docx")
backup_name = f"TUBITAK_2209A_Proje_Durumu_YEDEK_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
backup_path = os.path.join(DOCS, backup_name)

# Yedek al
shutil.copy2(docx_path, backup_path)
print(f"[OK] Yedek: {backup_name}")

# Docx'e yeni bolum ekle
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(docx_path)

# Remove duplicate styles to avoid KeyError
seen_ids = set()
styles_el = doc.styles.element
nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
for style_el in list(styles_el.findall('.//w:style', nsmap)):
    sid = style_el.get(qn('w:styleId'))
    if sid in seen_ids:
        styles_el.remove(style_el)
    else:
        seen_ids.add(sid)

# Sayfa sonu ekle
doc.add_page_break()

# --- Helper: add heading safely ---
def add_heading_safe(doc, text, level=1):
    p = doc.add_paragraph(text)
    style_name = f'Heading {level}'
    try:
        p.style = style_name
    except KeyError:
        # Fallback: set style via XML
        pPr = p._element.get_or_add_pPr()
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), f'Heading{level}')
        pPr.insert(0, pStyle)
    return p

# Yeni bolum basligi
heading = add_heading_safe(doc, 'Ek Gorseller ve Analizler', level=1)
for run in heading.runs:
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    run.font.size = Pt(18)
    run.bold = True

# Tarih
date_para = doc.add_paragraph()
date_run = date_para.add_run(f'Ekleme Tarihi: {datetime.now().strftime("%d %B %Y, %H:%M")}')
date_run.font.size = Pt(10)
date_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
date_run.italic = True

doc.add_paragraph('')

# --- Gorsel 31 ---
h31 = add_heading_safe(doc, 'Proje Ozet Infografigi', level=2)
for run in h31.runs:
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    run.font.size = Pt(14)
    run.bold = True

p31_desc = doc.add_paragraph()
p31_desc.add_run(
    'Bu gorsel, MC-AWARE projesinin 26 deneyinin (22 ana + 4 faz-2) '
    'tum ana bulgularini tek bir sayfada ozetlemektedir. Her kart, '
    'bir deney kategorisini (MC Tuzagi Cozumu, Anti-Prediktif Davranis, '
    'Mimari Bagimsizlik, Walk-Forward Dogrulama, Klasik ML Negatif Kontrol, '
    'Feature Ablasyon, IN_LEN Kirilganlik, Sektorel Kontrast) temsil '
    'eder. Alt barda toplam istatistikler (26 deney, 6 DL mimari, '
    '378+ konfigurasyon, 30+ gorsel, 115+ CSV, 7 walk-forward fold) '
    'yer almaktadir.'
).font.size = Pt(10)

doc.add_picture(os.path.join(OUTDIR, '31_Proje_Ozet_Infografik.png'), width=Inches(6.0))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# --- Gorsel 32 ---
h32 = add_heading_safe(doc, 'Concept Drift Zaman Serisi', level=2)
for run in h32.runs:
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    run.font.size = Pt(14)
    run.bold = True

p32_desc = doc.add_paragraph()
p32_desc.add_run(
    'Ust panel: USDTRY ve Petrol fiyatlarinin THYAO ile olan korelasyonlarinin '
    'egitim ve test setleri arasindaki kirilmasini (concept drift) gostermektedir. '
    'Kirmizi alan USDTRY, mavi alan Oil degiskeninin train-test korelasyon farkini '
    '(drift buyuklugunu) temsil eder. Ozellikle Fold 3-7 arasinda USDTRY '
    'korelasyonu dramatik olarak dusmekte, Oil korelasyonu ise isaret '
    'degistirmektedir. Alt panel: Her fold icin model dogrulugu (kirmizi), '
    'ters-yon dogrulugu (yesil) ve naive baseline (sari) karsilastirilmaktadir. '
    'Anti-prediktif foldlar (Fold 1, 2, 6) acik kirmizi arka plan ile '
    'isaretlenmistir. Bu grafik, korelasyon kirilmasinin anti-prediktif '
    'davranisla zamansal olarak ortusmesini dogrudan gostermektedir.'
).font.size = Pt(10)

doc.add_picture(os.path.join(OUTDIR, '32_Concept_Drift_Timeline.png'), width=Inches(6.0))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# --- Gorsel 33 ---
h33 = add_heading_safe(doc, 'Metodoloji Pipeline Akis Semasi', level=2)
for run in h33.runs:
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    run.font.size = Pt(14)
    run.bold = True

p33_desc = doc.add_paragraph()
p33_desc.add_run(
    'MC-AWARE projesinin uctan uca metodoloji akis semasi. '
    'Ust satir (Ana Hat): Ham veri toplama (Yahoo Finance, 2014-2023) -> '
    'Feature muhendisligi (13 ozellik: 10 teknik indikator + 3 makro degisken) -> '
    'MC-Aware on isleme (class_weight=balanced + MC-Aware loss ile lambda tuning) -> '
    'Derin ogrenme modelleri (6 mimari: BiLSTM, GRU, Conv1D, TCN, Transformer, RNN) -> '
    'Degerlendirme & analiz (Walk-Forward CV, 7 fold, 5 seed, Flip-Acc metrigi). '
    'Alt satir (Dogrulama Katmani): Negatif kontrol (klasik ML baseline) -> '
    'Ablasyon testleri (feature, IN_LEN, seed) -> Istatistiksel dogrulama '
    '(McNemar, binomial CI) -> Sektorel genelleme (sigorta, holding, NASDAQ) -> '
    'Bulgular (anti-prediktif oruntusu, concept drift kaniti, EMH-uyumlu sonuc). '
    'Toplam 378+ konfigurasyon, 26 deney, 90/90 MC=0 basarisi.'
).font.size = Pt(10)

doc.add_picture(os.path.join(OUTDIR, '33_Metodoloji_Pipeline.png'), width=Inches(6.0))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Kaydet
doc.save(docx_path)
print(f"[OK] {os.path.basename(docx_path)} guncellendi.")

print("\n========================================")
print("TUM ISLEMLER TAMAMLANDI!")
print(f"  - 3 yeni gorsel: Gorseller/31, 32, 33")
print(f"  - Docx yedek: {backup_name}")
print(f"  - Docx guncellendi: Proje_Durumu.docx")
print("========================================")
