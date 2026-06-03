"""
MC-AWARE Kisisel Rapor - Sifirdan Uretici
Profesyonel, gorseller ilgili bolumlerde, gercek akademik referanslarla
"""
import os
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

IMG = r'C:\Users\Kurt\Desktop\Tubitak\Gorseller'
OUT = r'C:\Users\Kurt\Desktop\Tubitak\Docs'

doc = Document()

# ════════════════════════════════════════════
# STYLES
# ════════════════════════════════════════════
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for lvl in [1, 2, 3]:
    hs = doc.styles['Heading {}'.format(lvl)]
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

# Helpers
def add_img(name, caption, width_cm=14):
    path = os.path.join(IMG, name)
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width_cm))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.size = Pt(9)
        r.italic = True
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()  # bosluk
    else:
        p = doc.add_paragraph()
        p.add_run('[Gorsel bulunamadi: {}]'.format(name)).italic = True

def add_para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def add_table_from_data(headers, rows):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Light Shading Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.rows[i+1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

# ════════════════════════════════════════════
# KAPAK SAYFASI
# ════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MC-AWARE')
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Anti-Prediktif Davranışın Derin Öğrenme ile Tespiti')
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TÜBİTAK 2209-A Üniversite Öğrencileri Araştırma Projeleri\nKişisel Araştırma Raporu')
r.font.size = Pt(12)

for _ in range(3):
    doc.add_paragraph()

info_lines = [
    ('Yürütücü:', 'Mehmet Ali Kurt'),
    ('Danışman:', 'Dr. Övgücan Karadağ Erdemir'),
    ('Üniversite:', 'Hacettepe Üniversitesi — Aktüerya Bilimleri'),
    ('Dönem:', 'Ekim 2025 – Haziran 2026'),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + ' ')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════
# ICINDEKILER (Placeholder)
# ════════════════════════════════════════════
doc.add_heading('İçindekiler', level=1)
add_para('(Word\'de bu tabloya sağ tıklayıp "Alanı Güncelle" seçiniz.)', italic=True, size=10)
doc.add_paragraph()

# Manual ToC
toc_items = [
    '1. Yönetici Özeti',
    '2. Giriş ve Motivasyon',
    '3. Literatür Taraması',
    '4. Metodoloji',
    '   4.1 Veri Kaynakları ve Özellik Mühendisliği',
    '   4.2 Model Mimarileri',
    '   4.3 MC-Aware Kayıp Fonksiyonu',
    '   4.4 Walk-Forward Doğrulama',
    '5. Deneysel Bulgular',
    '   5.1 MC Tuzağı Çözümü',
    '   5.2 Anti-Prediktif Davranış',
    '   5.3 Sektörel Ayrışma',
    '   5.4 Özellik Ablasyonu',
    '   5.5 Pencere Uzunluğu (IN_LEN) Ablasyonu',
    '   5.6 Walk-Forward Sonuçları',
    '   5.7 Ensemble ve Baseline',
    '   5.8 Çapraz Piyasa Analizi',
    '6. İstatistiksel Doğrulama',
    '7. Araştırmanın Kısıtları ve Gelecek Çalışmalar',
    '8. Sonuç',
    '9. Kaynakça',
]
for item in toc_items:
    add_para(item, size=10)

doc.add_page_break()

# ════════════════════════════════════════════
# BOLUM 1: YONETICI OZETI
# ════════════════════════════════════════════
doc.add_heading('1. Yönetici Özeti', level=1)

add_para(
    'Bu araştırma, TÜBİTAK 2209-A programı kapsamında gerçekleştirilen ve BIST (Borsa İstanbul) '
    'günlük yön tahmininde derin öğrenme modellerinin anti-prediktif davranışını sistematik olarak '
    'inceleyen kapsamlı bir çalışmadır.'
)

add_para(
    'Proje kapsamında 6 farklı derin öğrenme mimarisi (BiLSTM, GRU, Conv1D, TCN, Transformer, SimpleRNN), '
    '27 farklı varlık (11 BIST hissesi, 1 NASDAQ, 3 BES fonu, 12 sektörel test) ve 700+ konfigürasyon '
    'üzerinde 30+ deney serisi yürütülmüştür. Tüm deneyler 130 CSV dosyasında belgelenmiş olup, '
    '52 makale kalitesinde görsel üretilmiştir.'
)

add_para('Temel Bulgular:', bold=True)
add_para(
    '• MC Tuzağı: class_weight=balanced ile 700+ konfigürasyonda tamamen çözüldü.\n'
    '• Anti-Prediktif Davranış: 5/11 BIST hissesinde model, rastgeleden sistematik olarak kötü tahmin '
    'yapmakta; tahminlerin ters çevrilmesiyle doğruluk %60\'a ulaşmaktadır.\n'
    '• Mekanizma: Makroekonomik değişkenlerin (USD/TRY, petrol, faiz) eğitim-test dönemleri arasındaki '
    'korelasyon kırılması (USDTRY: eğitim r=0.91 → test r=0.41).\n'
    '• İstatistiksel Anlamlılık: Bonferroni-düzeltilmiş p = 0.00012.\n'
    '• Piyasaya Özgü: NASDAQ\'ta normal davranış — gelişmekte olan piyasalara özgü bir fenomen.'
)

add_img('31_Proje_Ozet_Infografik.png', 'Şekil 1: Proje Özet İnfografik')

# Proje olcegi tablosu
add_para('Proje Ölçeği:', bold=True)
add_table_from_data(
    ['Metrik', 'Değer'],
    [
        ['Toplam Konfigürasyon', '700+'],
        ['DL Mimarisi', '6 (BiLSTM, GRU, Conv1D, TCN, Transformer, SimpleRNN)'],
        ['Klasik ML Modeli', '4 (RF, DT, LR, SVM)'],
        ['Test Edilen Varlık', '27'],
        ['CSV Çıktı', '130 (71 özet + 19 tahmin + 26 diagnostik + 14 eşik)'],
        ['Görsel', '52'],
        ['Deney Serisi', '30+'],
        ['Walk-Forward', '7-fold v2 (3 seed)'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════
# BOLUM 2: GIRIS VE MOTIVASYON
# ════════════════════════════════════════════
doc.add_heading('2. Giriş ve Motivasyon', level=1)

add_para(
    'Finansal zaman serisi tahmini, makine öğrenmesi literatürünün en zorlu alanlarından biridir. '
    'Özellikle hisse senedi fiyat yönü tahmini, düşük sinyal-gürültü oranı (SNR), non-stasyonerlik '
    've piyasa verimsizliği gibi nedenlerle sınırlı başarı oranlarına sahiptir. Fama (1970) tarafından '
    'formüle edilen Etkin Piyasa Hipotezi (EMH), günlük fiyat hareketlerinin rastgele yürüyüş (random walk) '
    'izlediğini ve tahmin edilemez olduğunu öne sürmektedir.'
)

add_para(
    'Son yıllarda derin öğrenme (DL) modellerinin finansal tahmin alanında yaygınlaşmasıyla birlikte, '
    'araştırmacılar LSTM, GRU ve Transformer tabanlı mimarilerle %55-65 doğruluk oranları raporlamıştır '
    '(Fischer & Krauss, 2018; Sezer et al., 2020). Ancak bu çalışmaların büyük çoğunluğu gelişmiş piyasalar '
    '(S&P 500, NASDAQ) üzerine yoğunlaşmış olup, gelişmekte olan piyasalarda — özellikle Türkiye gibi yüksek '
    'volatiliteye sahip ekonomilerde — sonuçlar tutarsızdır.'
)

add_para(
    'Bu projenin ön çalışmasında (UYIK 2026 bildirisi), BIST\'te LSTM modellerinin Majority Class (MC) '
    'tuzağına düştüğü — yani modelin tüm tahminleri tek bir sınıfa yönlendirdiği — tespit edilmiştir. '
    'MC tuzağı çözüldükten sonra ise beklenmedik bir fenomen ortaya çıkmıştır: modeller rastgeleden '
    'sistematik olarak kötü performans sergilemektedir. Bu "anti-prediktif davranış", projenin temel '
    'araştırma sorusunu oluşturmuştur.'
)

add_para('Araştırma Soruları:', bold=True)
add_para(
    '1. MC tuzağı çözüldükten sonra DL modelleri neden rastgeleden kötü performans sergiler?\n'
    '2. Bu davranış mimariden bağımsız mı, yoksa belirli mimarilere özgü mü?\n'
    '3. Hangi piyasa koşulları ve hangi hisseler anti-prediktif davranışa yatkındır?\n'
    '4. Anti-prediktif davranışın mekanizması nedir ve nasıl önlenebilir?'
)

add_img('33_Metodoloji_Pipeline.png', 'Şekil 2: Proje Metodoloji Pipeline')

doc.add_page_break()

# ════════════════════════════════════════════
# BOLUM 3: LITERATUR TARAMASI
# ════════════════════════════════════════════
doc.add_heading('3. Literatür Taraması', level=1)

# 3.1
doc.add_heading('3.1 Finansal Zaman Serilerinde Derin Öğrenme', level=2)
add_para(
    'Fischer ve Krauss (2018), S&P 500 üzerinde LSTM ağlarının Random Forest ve Logistic Regression\'a '
    'göre daha yüksek doğruluk sağladığını göstermiştir. Ancak yazarlar, derin öğrenmenin gürültülü '
    'piyasa ortamlarında aşırı uyuma (overfitting) eğilimli olduğunu da vurgulamıştır. Sezer, Gudelek ve '
    'Ozbayoglu (2020) kapsamlı bir sistematik derleme çalışmasında, finansal tahmin için kullanılan 150+ '
    'DL çalışmasını inceleyerek, CNN ve LSTM tabanlı yaklaşımların en yaygın mimariler olduğunu '
    'raporlamıştır.'
)
add_para(
    'Referans: Fischer, T., & Krauss, C. (2018). Deep learning with long short-term memory networks '
    'for financial market predictions. European Journal of Operational Research, 270(2), 654-669.', italic=True, size=10
)
add_para(
    'Referans: Sezer, O.B., Gudelek, M.U., & Ozbayoglu, A.M. (2020). Financial time series forecasting '
    'with deep learning: A systematic literature review. Applied Soft Computing, 90, 106181.', italic=True, size=10
)

# 3.2
doc.add_heading('3.2 Türkiye\'de BIST Üzerine Derin Öğrenme Çalışmaları', level=2)
add_para(
    'Türkiye\'de BIST üzerine yapılan DL çalışmaları sınırlıdır. Kara, Boyacıoğlu ve Baykan (2011), '
    'İMKB-100 endeksi üzerinde yapay sinir ağlarının (ANN) geleneksel istatistiksel yöntemlerden daha '
    'iyi performans gösterdiğini raporlamıştır. Göçken et al. (2016), BIST-100 için hibrit ANN-metaheuristik '
    'yaklaşımlar önermiştir. Ancak bu çalışmaların hiçbiri anti-prediktif davranışı incelememiş veya '
    'MC tuzağını ele almamıştır.'
)
add_para(
    'Referans: Kara, Y., Boyacıoğlu, M.A., & Baykan, Ö.K. (2011). Predicting direction of stock price '
    'index movement using artificial neural networks. Expert Systems with Applications, 38(5), 5311-5319.', italic=True, size=10
)

# 3.3
doc.add_heading('3.3 Sınıf Dengesizliği ve MC Tuzağı', level=2)
add_para(
    'Finansal tahmin problemlerinde sınıf dengesizliği kritik bir sorundur. He ve Garcia (2009), '
    'dengesiz veri setlerinde sınıflandırıcıların çoğunluk sınıfına yönelme eğilimini kapsamlı olarak '
    'incelemiştir. Krawczyk (2016), cost-sensitive learning yöntemlerinin bu soruna etkili çözümler '
    'sunduğunu göstermiştir. Projemizde class_weight=balanced parametresi ile BCE loss kombinasyonu, '
    '700+ konfigürasyonun tamamında MC tuzağını ortadan kaldırmıştır.'
)
add_para(
    'Referans: He, H., & Garcia, E.A. (2009). Learning from imbalanced data. IEEE Transactions on '
    'Knowledge and Data Engineering, 21(9), 1263-1284.', italic=True, size=10
)
add_para(
    'Referans: Krawczyk, B. (2016). Learning from imbalanced data: open challenges and future '
    'directions. Progress in Artificial Intelligence, 5(4), 221-232.', italic=True, size=10
)

# 3.4
doc.add_heading('3.4 Kısayol Öğrenme ve Sahte Korelasyonlar', level=2)
add_para(
    'Geirhos et al. (2020), derin öğrenme modellerinin eğitim verisindeki sahte korelasyonlara '
    '(spurious correlations) dayanarak "kısayol" çözümler öğrendiğini göstermiştir. Arjovsky et al. (2019), '
    'bu tür sahte korelasyonların dağılım kayması (distribution shift) altında model performansını '
    'dramatik biçimde düşürdüğünü kanıtlamıştır. Projemizdeki anti-prediktif davranış, bu teorik '
    'çerçevenin finansal piyasalardaki ampirik kanıtıdır: modeller eğitim dönemindeki makro '
    'korelasyonları öğrenmekte, ancak test döneminde bu korelasyonlar kırıldığında sistematik olarak '
    'yanlış tahmin üretmektedir.'
)
add_para(
    'Referans: Geirhos, R., et al. (2020). Shortcut learning in deep neural networks. Nature Machine '
    'Intelligence, 2(11), 665-673.', italic=True, size=10
)
add_para(
    'Referans: Arjovsky, M., et al. (2019). Invariant risk minimization. arXiv preprint arXiv:1907.02893.', italic=True, size=10
)

# 3.5
doc.add_heading('3.5 Kavram Kayması ve Piyasa Rejim Değişikliği', level=2)
add_para(
    'Gama et al. (2014), makine öğrenmesinde kavram kaymasını (concept drift) — veri dağılımının '
    'zamanla değişmesi — kapsamlı olarak incelemiştir. Finansal piyasalarda kavram kayması, '
    'rejim değişiklikleri (bull/bear, kriz dönemleri) olarak tezahür eder. Bekaert ve Harvey (2002), '
    'gelişmekte olan piyasaların gelişmiş piyasalara kıyasla daha sık ve şiddetli rejim değişikliklerine '
    'maruz kaldığını göstermiştir. Projemizdeki Walk-Forward analizi, Fold 7\'nin (2024-2025 enflasyon '
    'dönemi) en yoğun anti-prediktif davranışı sergilediğini ortaya koymuştur — bu, kavram kaymasının '
    'doğrudan kanıtıdır.'
)
add_para(
    'Referans: Gama, J., et al. (2014). A survey on concept drift adaptation. ACM Computing Surveys, 46(4), 1-37.', italic=True, size=10
)
add_para(
    'Referans: Bekaert, G., & Harvey, C.R. (2002). Research in emerging markets finance. '
    'Journal of Empirical Finance, 9(1), 3-55.', italic=True, size=10
)

# 3.6
doc.add_heading('3.6 Etkin Piyasa Hipotezi ve Anti-Prediktif Davranış', level=2)
add_para(
    'Fama (1970), zayıf formda piyasa etkinliğinin geçmiş fiyat bilgisiyle gelecek fiyat hareketlerinin '
    'tahmin edilemeyeceğini savunur. Projemizdeki bulgular EMH\'yi genişletmektedir: modeller sadece '
    'başarısız olmakla kalmayıp, rastgeleden sistematik olarak kötü performans sergilemektedir. Bu, '
    'De Long et al. (1990) tarafından tanımlanan "noise trader risk" kavramıyla tutarlıdır — spekülatif '
    'piyasalarda gürültü tüccarlarının davranışı, rasyonel modelleri yanıltmaktadır.'
)
add_para(
    'Referans: Fama, E.F. (1970). Efficient capital markets: A review of theory and empirical work. '
    'The Journal of Finance, 25(2), 383-417.', italic=True, size=10
)
add_para(
    'Referans: De Long, J.B., et al. (1990). Noise trader risk in financial markets. '
    'Journal of Political Economy, 98(4), 703-738.', italic=True, size=10
)

# 3.7
doc.add_heading('3.7 Walk-Forward Doğrulama', level=2)
add_para(
    'Tashman (2000), zaman serisi tahmininde walk-forward doğrulamanın standart k-fold cross-validation\'a '
    'göre daha güvenilir sonuçlar verdiğini göstermiştir. Cerqueira, Torgo ve Mozetič (2020), '
    'walk-forward\'ın veri sızıntısını önlemedeki kritik rolünü vurgulamıştır. Projemizde 7-fold walk-forward '
    'v2 yaklaşımı, val_data düzeltmesi ile uygulanmış ve her fold 6 mimari × 3 seed ile test edilmiştir.'
)
add_para(
    'Referans: Tashman, L.J. (2000). Out-of-sample tests of forecasting accuracy. '
    'International Journal of Forecasting, 16(4), 437-450.', italic=True, size=10
)

# 3.8
doc.add_heading('3.8 Seçici Tahmin ve Model Güvenilirliği', level=2)
add_para(
    'Geifman ve El-Yaniv (2017), "selective prediction" kavramını tanımlayarak, modelin düşük güvenirlikli '
    'tahminlerde karar vermekten kaçınmasının genel doğruluğu artırdığını kanıtlamıştır. '
    'Bu yaklaşım, projemizin gelecek çalışma paketlerinden biri olarak belirlenmiştir: anti-prediktif '
    'olduğu tespit edilen dönemlerde modelin "abstain" (kaçınma) kararı vermesi.'
)
add_para(
    'Referans: Geifman, Y., & El-Yaniv, R. (2017). Selective classification for deep neural networks. '
    'Advances in Neural Information Processing Systems (NeurIPS), 30.', italic=True, size=10
)

doc.add_page_break()

# ════════════════════════════════════════════
# BOLUM 4: METODOLOJI
# ════════════════════════════════════════════
doc.add_heading('4. Metodoloji', level=1)

doc.add_heading('4.1 Veri Kaynakları ve Özellik Mühendisliği', level=2)
add_para(
    'Veri kaynakları: Yahoo Finance (OHLCV), TCMB (faiz oranları), FRED (WTI petrol fiyatları), '
    'CNN Business (Fear & Greed Index). Tüm veriler günlük frekansta, 2019-2025 dönemini kapsamaktadır.'
)
add_para('13 Özellik Vektörü:', bold=True)
add_table_from_data(
    ['#', 'Özellik', 'Kaynak', 'Tür'],
    [
        ['1-5', 'Open, High, Low, Close, Volume', 'Yahoo Finance', 'Fiyat'],
        ['6', 'SMA_10 / Close oranı', 'Hesaplanan', 'Teknik'],
        ['7', 'EMA_10 / Close oranı', 'Hesaplanan', 'Teknik'],
        ['8', 'RSI_14', 'Hesaplanan', 'Teknik'],
        ['9', 'Log Return', 'Hesaplanan', 'İstatistik'],
        ['10', 'SMA_Cross (5/20)', 'Hesaplanan', 'Teknik'],
        ['11', 'USD/TRY kuru', 'Yahoo Finance', 'Makro'],
        ['12', 'WTI Petrol fiyatı', 'FRED', 'Makro'],
        ['13', 'TCMB Politika Faizi', 'TCMB', 'Makro'],
    ]
)

add_img('39_Label_Distribution.png', 'Şekil 3: Hedef Değişken Dağılımı (Sınıf Dengesi)')

doc.add_heading('4.2 Model Mimarileri', level=2)
add_para(
    '6 derin öğrenme mimarisi test edilmiştir. Tüm modeller aynı hiperparametre uzayında '
    '(epochs=100, batch_size=32, learning_rate=0.001, early_stopping patience=10, dropout=0.3) '
    'eğitilmiştir.'
)
add_table_from_data(
    ['Mimari', 'Açıklama', 'Parametre Sayısı'],
    [
        ['BiLSTM', 'Bidirectional LSTM (64+32 unit)', '~25K'],
        ['GRU', 'Gated Recurrent Unit (64+32 unit)', '~19K'],
        ['Conv1D', '1D Convolutional (64 filter, kernel=3)', '~15K'],
        ['TCN', 'Temporal Convolutional Network (dilation 1,2,4)', '~20K'],
        ['Transformer', 'Multi-Head Attention (4 head, d=64)', '~30K'],
        ['SimpleRNN', 'Vanilla RNN (64+32 unit)', '~12K'],
    ]
)

add_img('06_Mimari_Kiyaslama.png', 'Şekil 4: 6 DL Mimarisi Karşılaştırması — Model vs Flip vs Naive')

doc.add_heading('4.3 MC-Aware Kayıp Fonksiyonu', level=2)
add_para(
    'Standart Binary Cross-Entropy (BCE) loss\'a ek olarak, MC tuzağını tespit eden bir penaltı terimi '
    'eklenmiştir. class_weight=balanced parametresi ile kombinlenerek, modelin tek sınıfa yönelmesi '
    'engellenmiştir. λ (lambda) hiperparametresi ile MC penaltisinin ağırlığı kontrol edilmektedir.'
)

add_img('05_MC_Tuzagi_Cozumu.png', 'Şekil 5: MC Tuzağı Çözümü — class_weight + MC-Aware Loss')

doc.add_heading('4.4 Walk-Forward Doğrulama', level=2)
add_para(
    'Finansal veride standart k-fold cross-validation veri sızıntısına (data leakage) neden olur '
    'çünkü gelecek veri eğitimde kullanılabilir. Walk-Forward yaklaşımında eğitim penceresi her fold\'da '
    'ileriye kayar ve model asla geleceği görmez. Projede 7-fold walk-forward v2 kullanılmıştır: '
    'her fold 6 mimari × 3 seed = 18 deney içerir (toplam 126 deney).'
)

add_img('40_WalkForward_v2_Heatmap.png',
        'Şekil 6: Walk-Forward v2 — Fold × Mimari Doğruluk Haritası (3 Seed Ortalaması)')

doc.add_page_break()

# ════════════════════════════════════════════
# BOLUM 5: DENEYSEL BULGULAR
# ════════════════════════════════════════════
doc.add_heading('5. Deneysel Bulgular', level=1)

# 5.1
doc.add_heading('5.1 MC Tuzağı Çözümü', level=2)
add_para(
    'Tüm 700+ konfigürasyonda class_weight=balanced + BCE loss kombinasyonu ile MC tuzağı tamamen '
    'çözülmüştür. Hiçbir konfigürasyonda model tek sınıfa yönelmemiştir (MC_count = 0).'
)

# 5.2
doc.add_heading('5.2 Anti-Prediktif Davranış', level=2)
add_para(
    'MC tuzağı çözüldükten sonra ortaya çıkan temel bulgu: DL modelleri BIST\'te sadece başarısız '
    'olmamakta, sistematik olarak yanlış yönü tahmin etmektedir. full_13 özellik grubuyla eğitilen '
    'modellerin ortalama doğruluğu %39.6 iken, tahminlerin ters çevrilmesiyle (flip accuracy) %60.4\'e '
    'ulaşılmaktadır. Bu fark, Bonferroni-düzeltilmiş p = 0.00012 ile istatistiksel olarak anlamlıdır.'
)

add_img('02_Anti_Prediktif_Scatter.png', 'Şekil 7: Anti-Prediktif Davranış — Model Acc vs Flip Acc')
add_img('G3_anti_prediktif_heatmap.png', 'Şekil 8: Anti-Prediktif Davranış Isı Haritası')

# 5.3
doc.add_heading('5.3 Sektörel Ayrışma', level=2)
add_para(
    '11 BIST hissesi üzerinde yapılan kapsamlı test, anti-prediktif davranışın sektörel bir örüntü '
    'izlediğini ortaya koymuştur. Havacılık (THYAO, PGSUS) ve spekülatif (SASA, HEKTS) hisselerinde '
    '%100 anti-prediktif davranış gözlenirken, bankacılık (YKBNK) ve otomotiv (FROTO) sektörlerinde '
    'normal davranış sergilenmektedir.'
)

add_table_from_data(
    ['Hisse', 'Sektör', 'Flip Wins', 'Oran', 'Durum'],
    [
        ['THYAO', 'Havacılık', '15/15', '%100', '🔴 Anti-Prediktif'],
        ['PGSUS', 'Havacılık', '15/15', '%100', '🔴 Anti-Prediktif'],
        ['HEKTS', 'Kimya/Spekülatif', '15/15', '%100', '🔴 Anti-Prediktif'],
        ['SASA', 'Kimya/Spekülatif', '15/15', '%100', '🔴 Anti-Prediktif'],
        ['KRDMD', 'Demir-Çelik', '13/15', '%87', '🟠 Güçlü'],
        ['DOAS', 'Otomotiv/Ticaret', '10/15', '%67', '🟡 Kısmi'],
        ['ISCTR', 'Bankacılık', '7/15', '%47', '⚪ Nötr'],
        ['TAVHL', 'Havalimanı', '7/15', '%47', '⚪ Nötr'],
        ['FROTO', 'Otomotiv', '1/15', '%7', '🟢 Normal'],
        ['YKBNK', 'Bankacılık', '0/15', '%0', '🟢 Normal'],
    ]
)

add_img('34_PerStock_AccBar_11BIST.png',
        'Şekil 9: 11 BIST Hissesi — Naive, Model ve Flip Doğruluk Karşılaştırması')
add_img('01_Sektorel_Kiyaslama_Bar.png', 'Şekil 10: Sektörel Karşılaştırma')
add_img('41_SeedVariance_AllStocks.png', 'Şekil 11: Seed Varyansı — 11 BIST Hissesi')

# 5.4
doc.add_heading('5.4 Özellik Ablasyonu', level=2)
add_para(
    'Anti-prediktif davranışın kaynağını belirlemek için özellik ablasyonu yapılmıştır. 13 özellik '
    'iki gruba ayrılmıştır: full_13 (tüm özellikler) ve no_ext_10 (makro değişkenler — USD/TRY, '
    'petrol, TCMB faizi — çıkarılmış). Sonuçlar çarpıcıdır:'
)

add_table_from_data(
    ['Grup', 'Özellik Sayısı', 'Ortalama Acc', 'Flip Acc', 'Anti-Prediktif?'],
    [
        ['full_13', '13', '%39.6', '%60.4', 'EVET'],
        ['no_ext_10', '10', '%49.6', '%50.4', 'HAYIR'],
    ]
)

add_para(
    'Makro değişkenler çıkarıldığında anti-prediktif davranış tamamen ortadan kalkmaktadır. Bu, '
    'anti-prediktif davranışın kaynağının makroekonomik korelasyon kırılması olduğunu kanıtlamaktadır.'
)

add_img('04_Feature_Ablation_Bar.png', 'Şekil 12: Özellik Ablasyonu — full_13 vs no_ext_10')
add_img('03_Correlation_Drift_Slope.png', 'Şekil 13: Korelasyon Kayması — Eğitim vs Test')
add_img('42_CorrelationDrift_Detail.png', 'Şekil 14: Korelasyon Kayması Detay')

# 5.5
doc.add_heading('5.5 Pencere Uzunluğu (IN_LEN) Ablasyonu', level=2)
add_para(
    'Giriş penceresinin uzunluğu (IN_LEN) anti-prediktif davranışı doğrudan etkilemektedir. '
    'IN_LEN=2 ve IN_LEN=5 ile modeller anti-prediktif davranış sergilerken, IN_LEN=10 ile '
    'davranış normalleşmektedir. Bu, kısa pencerelerin makro gürültüyü yakalamaya daha yatkın '
    'olduğunu göstermektedir.'
)

add_table_from_data(
    ['IN_LEN', 'Ortalama Acc', 'Anti-Prediktif?'],
    [
        ['2', '%48.5', 'Evet — tetikleniyor'],
        ['5', '%51.5', 'Sınırda'],
        ['10', '%52.2', 'Hayır — normal'],
    ]
)

add_img('48_INLEN_Ablation_Detail.png',
        'Şekil 15: IN_LEN Ablasyonu — Pencere Uzunluğu Etkisi ve Anti-Prediktif Sayısı')
add_img('15_IN_LEN_Ablasyonu.png', 'Şekil 16: IN_LEN Ablasyonu Genel Görünüm')

# 5.6
doc.add_heading('5.6 Walk-Forward Sonuçları', level=2)
add_para(
    '7-fold walk-forward v2 sonuçları, anti-prediktif davranışın dönemsel bir örüntü izlediğini '
    'göstermektedir. Fold 7 (2024-2025 yüksek enflasyon dönemi) en yoğun anti-prediktif oranına '
    'sahiptir (8/18). Fold 1, 3, 4, 5\'te hiç anti-prediktif davranış yoktur.'
)

add_table_from_data(
    ['Fold', 'Anti-Prediktif', 'Toplam', 'Oran', 'Dönem'],
    [
        ['1', '0', '18', '%0', 'Normal piyasa'],
        ['2', '2', '18', '%11', 'COVID şoku'],
        ['3', '0', '18', '%0', 'Toparlanma'],
        ['4', '0', '18', '%0', 'Normal'],
        ['5', '0', '18', '%0', 'Normal'],
        ['6', '1', '18', '%6', 'Kısmi'],
        ['7', '8', '18', '%44', 'Enflasyon dönemi'],
    ]
)

add_img('16_WalkForward_MultiArch_Heatmap.png',
        'Şekil 17: Walk-Forward Mimari × Fold Isı Haritası')
add_img('44_PooledConfusion_PerFold.png',
        'Şekil 18: Walk-Forward Fold Bazlı Confusion Metrikleri')
add_img('30_WalkForward_FoldByFold.png', 'Şekil 19: Walk-Forward Fold Detayları')

# 5.7
doc.add_heading('5.7 Ensemble ve Baseline Karşılaştırması', level=2)
add_para(
    'Hard Voting (%39.7) ve Soft Voting (%37.7) yöntemleri anti-prediktif davranışı düzeltmemiştir. '
    'Bireysel mimariler arasında SimpleRNN (%45.6) en yüksek doğruluğa sahip olsa da, hiçbir DL '
    'mimarisi naive baseline\'ı tutarlı biçimde geçememiştir. Klasik ML modelleri (RF, DT, LR, SVM) '
    'ise anti-prediktif davranış sergilememektedir — bu sorun yalnızca DL\'ye özgüdür.'
)

add_img('37_Ensemble_Agreement.png', 'Şekil 20: Ensemble Uzlaşma Analizi')
add_img('19_Majority_Voting_vs_ML.png', 'Şekil 21: Majority Voting vs Klasik ML')
add_img('25_McNemar_Matrix.png', 'Şekil 22: McNemar Test Matrisi — Mimariler Arası İstatistiksel Fark')

# 5.8
doc.add_heading('5.8 Çapraz Piyasa Analizi', level=2)
add_para(
    'NASDAQ (AAPL) üzerinde aynı pipeline uygulandığında anti-prediktif davranış gözlenmemiştir. '
    'Bu, fenomenin gelişmekte olan piyasalara — özellikle Türkiye\'nin yüksek makroekonomik volatilitesine — '
    'özgü olduğunu göstermektedir.'
)

add_img('38_CrossMarket_Ablation.png', 'Şekil 23: Çapraz Piyasa Özellik Ablasyonu — NASDAQ vs BIST')
add_img('28_AAPL_vs_THYAO_GroupCompare.png', 'Şekil 24: AAPL vs THYAO Karşılaştırma')

# Ek gorseller
doc.add_heading('5.9 Sektörel Detay Analizleri', level=2)
add_para('Sigorta ve holding sektörlerinde ayrıntılı seed varyansı ve BES fonları analizi:')

add_img('24_Sigorta_vs_Holding_PerStock.png', 'Şekil 25: Sigorta vs Holding Sektörü Per-Stock')
add_img('46_Sigorta_SeedVar.png', 'Şekil 26: Sigorta Sektörü Seed Varyansı')
add_img('45_Holding_SeedVar.png', 'Şekil 27: Holding Sektörü Seed Varyansı')
add_img('43_BES_Fund_Detail.png', 'Şekil 28: BES Fonları Karşılaştırma')

doc.add_page_break()

# ════════════════════════════════════════════
# BOLUM 6: ISTATISTIKSEL DOGRULAMA
# ════════════════════════════════════════════
doc.add_heading('6. İstatistiksel Doğrulama', level=1)

doc.add_heading('6.1 Binom Testi', level=2)
add_para(
    'Anti-prediktif davranışın istatistiksel anlamlılığı binom testi ile doğrulanmıştır. '
    'H₀: Model doğruluğu rastgele (%50). THYAO\'da 15/15 konfigürasyonda flip > naive '
    '(p < 0.0001). Bonferroni düzeltmesi ile çoklu karşılaştırma dikkate alınmış, '
    'düzeltilmiş p = 0.00012 elde edilmiştir.'
)

doc.add_heading('6.2 McNemar Testi', level=2)
add_para(
    '6 DL mimarisinin tahminleri arasındaki istatistiksel fark McNemar testi ile incelenmiştir. '
    'Sonuçlar, mimarilerin benzer hata kalıplarına sahip olduğunu — yani anti-prediktif davranışın '
    'mimariden bağımsız olduğunu — göstermiştir.'
)

doc.add_heading('6.3 Karşılıklı Bilgi (MI) Analizi', level=2)
add_para(
    'Özellikler ile hedef değişken arasındaki karşılıklı bilgi (mutual information) miktarı '
    'ölçülmüştür. Düşük MI skorları, özelliklerin hedef değişkenle sınırlı bilgi-teorik bağlantıya '
    'sahip olduğunu doğrulamaktadır.'
)

add_img('36_MI_Scores_Bar.png', 'Şekil 29: Karşılıklı Bilgi (MI) Skorları')
add_img('09_AntiPredictive_ROC_Curve.png', 'Şekil 30: Anti-Prediktif ROC Eğrisi')
add_img('20_Pooled_Confusion_Matrix.png', 'Şekil 31: Havuzlanmış Confusion Matrix')

doc.add_heading('6.4 Ek Diagnostik Görseller', level=2)
add_img('35_Threshold_Heatmap.png', 'Şekil 32: Eşik Optimizasyonu Yüzey Haritası')
add_img('47_Threshold_MultiArch.png', 'Şekil 33: Mimari Bazlı Eşik-Doğruluk Eğrileri')
add_img('23_Lambda_Acc_FlipAcc.png', 'Şekil 34: Lambda vs Accuracy/Flip Accuracy')
add_img('13_Seed_Invariance_Bulgusu.png', 'Şekil 35: Seed İnvaryansı Bulgusu')

doc.add_page_break()

# ════════════════════════════════════════════
# BOLUM 7: KISITLAR
# ════════════════════════════════════════════
doc.add_heading('7. Araştırmanın Kısıtları ve Gelecek Çalışmalar', level=1)

limitations = [
    ('Spekülatif Rejimlerin Dinamik Tespiti',
     'Hangi dönemlerde anti-prediktif davranışın tetikleneceği önceden tahmin edilememektedir.',
     'Markov Regime-Switching modelleri ile rejim tespiti.'),
    ('Statik Eşik Değerleri',
     'Olasılık eşiği tüm dönemler için sabit (0.50) kullanılmıştır.',
     'Selective Abstain yaklaşımı (Geifman & El-Yaniv, 2017).'),
    ('MI Bias',
     'Karşılıklı bilgi tahmincileri sonlu örneklem yanlılığına sahiptir.',
     'kNN-tabanlı MI tahmincileri.'),
    ('Zaman Frekansı Kısıtı',
     'Günlük frekans yüksek gürültü içerir. IN_LEN ablasyonu kısmen çözmüştür.',
     'Haftalık ve aylık periyotlarla test.'),
    ('Donanım Kısıtları',
     'Tüm deneyler tek GPU üzerinde yapılmıştır.',
     'Bulut bilişim ile daha geniş hiperparametre arama.'),
    ('Transfer Learning Eksikliği',
     'Gelişmiş piyasalardan (NASDAQ) BIST\'e transfer öğrenme denenmemiştir.',
     'Domain adaptation teknikleri.'),
    ('Çoklu Karşılaştırma Riski',
     '700+ konfigürasyon test edilmesi çoklu karşılaştırma riskini artırır.',
     'Bonferroni düzeltmesi uygulanmıştır (p = 0.00012).'),
    ('Validasyon Metodolojisi',
     'Walk-forward v1\'de val_data hatası tespit edilip v2\'de düzeltilmiştir.',
     'v2 sonuçları referans alınmaktadır.'),
    ('Makro Değişken Gecikmesi',
     'Makro değişkenler aynı gün kullanılmış, gecikme analizi yapılmamıştır.',
     'Lag analizi ile zamanlama optimizasyonu.'),
    ('Tek Piyasa Kısıtı',
     'Sadece BIST ve kısmen NASDAQ test edilmiştir.',
     'Diğer gelişmekte olan piyasalar (BSE India, JSE South Africa).'),
    ('Duygu Analizi Eksikliği',
     'Sosyal medya ve haber duygu analizi dahil edilmemiştir.',
     'Twitter/X sentiment + NLP entegrasyonu.'),
    ('İkili Etiketleme Basitleştirmesi',
     'Hedef değişken sadece yükseldi/düştü olarak tanımlanmıştır.',
     'Üçlü etiketleme (yükseldi/düştü/nötr) ve regresyon.'),
    ('İşlem Maliyeti Analizi',
     'Komisyon, spread ve slippage dahil edilmemiştir.',
     'Gerçekçi portföy simülasyonu.'),
    ('Hayatta Kalma Yanlılığı',
     'Sadece hâlâ borsada işlem gören hisseler test edilmiştir.',
     'Borsadan çıkan hisselerin dahil edilmesi.'),
    ('Açıklanabilirlik',
     'DL modelleri "kara kutu" niteliğindedir.',
     'SHAP, LIME, Grad-CAM ile açıklanabilirlik analizi.'),
]

for i, (title, desc, future) in enumerate(limitations, 1):
    doc.add_heading('Kısıt {}: {}'.format(i, title), level=2)
    add_para(desc)
    p = doc.add_paragraph()
    r = p.add_run('Gelecek Çalışma: ')
    r.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run(future)
    r2.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════
# BOLUM 8: SONUC
# ════════════════════════════════════════════
doc.add_heading('8. Sonuç', level=1)

add_para(
    'Bu araştırma, TÜBİTAK 2209-A programı kapsamında lisans düzeyinde gerçekleştirilmiş '
    'kapsamlı bir çalışmadır. 6 DL mimarisi, 27 varlık, 700+ konfigürasyon ve 130 CSV çıktısı ile '
    'BIST\'te derin öğrenme modellerinin anti-prediktif davranışı sistematik olarak incelenmiştir.'
)

add_para('Projenin Temel Katkıları:', bold=True)
add_para(
    '1. MC Tuzağı Çözümü: class_weight=balanced ile 700+ konfigürasyonda MC=0 sağlanmıştır.\n'
    '2. Anti-Prediktif Davranışın Tespiti: 5/11 BIST hissesinde sistematik ters tahmin.\n'
    '3. Mekanizma Analizi: Makro korelasyon kırılması mekanizması kanıtlanmıştır.\n'
    '4. Sektörel Harita: Havacılık ve spekülatif sektörler risk altında.\n'
    '5. Mimariden Bağımsızlık: 6 DL mimarisinde aynı örüntü.\n'
    '6. Pencere Uzunluğu Etkisi: IN_LEN ≤ 5 anti-prediktif tetikler, IN_LEN=10 normalleştirir.\n'
    '7. Piyasaya Özgülük: NASDAQ\'ta normal davranış — gelişmekte olan piyasalara özgü.'
)

add_para(
    'Projenin pratik implikasyonu açıktır: BIST\'te DL tabanlı yön tahmini modellerinin doğrudan '
    'kullanımı, belirli sektörlerde ve dönemlerde yatırımcıya zarar verebilir. Bu çalışma, '
    '"teşhis" aşamasını tamamlamıştır. Gelecek çalışmalar, Selective Abstain ve Regime-Switching '
    'yöntemleriyle "tedavi" aşamasına geçmeyi hedeflemektedir.'
)

add_para(
    'Araştırmanın 15 kısıtı ve 15 gelecek çalışma iş paketi belirlenmiştir. '
    'Tüm veriler, kodlar ve sonuçlar açık kaynak olarak GitHub\'da paylaşılmıştır.',
)

doc.add_page_break()

# ════════════════════════════════════════════
# BOLUM 9: KAYNAKCA
# ════════════════════════════════════════════
doc.add_heading('9. Kaynakça', level=1)

add_para(
    'Aşağıdaki kaynaklar APA 7 formatında listelenmiştir. Her referansın altında '
    'DOI/URL bağlantısı ve projemizle ilişkisi açıklanmıştır.',
    italic=True, size=10
)

refs = [
    ('Arjovsky, M., Bottou, L., Gulrajani, I., & Lopez-Paz, D. (2019). Invariant risk minimization. arXiv preprint arXiv:1907.02893.',
     'https://arxiv.org/abs/1907.02893',
     'Sahte korelasyonların dağılım kayması altında model performansını düşürdüğünü kanıtlar. Projemizdeki makro korelasyon kırılması bu teorinin ampirik kanıtıdır.'),

    ('Bekaert, G., & Harvey, C.R. (2002). Research in emerging markets finance: Looking ahead. Journal of Empirical Finance, 9(1), 3-55.',
     'https://doi.org/10.1016/S0927-5398(01)00043-2',
     'Gelişmekte olan piyasaların gelişmiş piyasalara kıyasla daha kırılgan olduğunu gösterir. BIST\'te anti-prediktif davranışın NASDAQ\'ta görülmemesi bu tezle tutarlıdır.'),

    ('Cerqueira, V., Torgo, L., & Mozetič, I. (2020). Evaluating time series forecasting models: An empirical study. Machine Learning, 109(11), 1997-2028.',
     'https://doi.org/10.1007/s10994-020-05910-7',
     'Walk-forward validasyonun zaman serilerinde standart k-fold\'dan daha güvenilir olduğunu kanıtlar. Projemizdeki 7-fold WF v2 yaklaşımının metodolojik temelini oluşturur.'),

    ('De Long, J.B., Shleifer, A., Summers, L.H., & Waldmann, R.J. (1990). Noise trader risk in financial markets. Journal of Political Economy, 98(4), 703-738.',
     'https://doi.org/10.1086/261703',
     'Gürültü tüccarlarının rasyonel modelleri yanılttığını teorize eder. SASA ve HEKTS gibi spekülatif hisselerdeki %100 anti-prediktif davranış, noise trader risk\'in doğrudan kanıtıdır.'),

    ('Fama, E.F. (1970). Efficient capital markets: A review of theory and empirical work. The Journal of Finance, 25(2), 383-417.',
     'https://doi.org/10.2307/2325486',
     'Etkin Piyasa Hipotezi\'nin temel kaynağı. Projemiz EMH\'yi genişletir: modeller sadece başarısız olmayıp, sistematik olarak yanlış tahmin üretmektedir.'),

    ('Fischer, T., & Krauss, C. (2018). Deep learning with long short-term memory networks for financial market predictions. European Journal of Operational Research, 270(2), 654-669.',
     'https://doi.org/10.1016/j.ejor.2017.11.054',
     'S&P 500\'de LSTM\'in klasik yöntemlerden üstün olduğunu gösterir. Projemiz bunun gelişmekte olan piyasalarda geçerli olmadığını — hatta ters etki yaptığını — kanıtlar.'),

    ('Gama, J., Žliobaitė, I., Bifet, A., Pechenizkiy, M., & Bouchachia, A. (2014). A survey on concept drift adaptation. ACM Computing Surveys, 46(4), 1-37.',
     'https://doi.org/10.1145/2523813',
     'Kavram kaymasının (concept drift) kapsamlı tanımını yapar. Walk-Forward Fold 7\'deki ani anti-prediktif artış, kavram kaymasının doğrudan kanıtıdır.'),

    ('Geifman, Y., & El-Yaniv, R. (2017). Selective classification for deep neural networks. Advances in Neural Information Processing Systems (NeurIPS), 30.',
     'https://proceedings.neurips.cc/paper/2017/hash/4a8423d5e91fda00bb7e46540e2b0cf1-Abstract.html',
     'Modelin düşük güvenirlikli tahminlerde kaçınmasının doğruluğu artırdığını kanıtlar. Projemizin gelecek çalışma paketi: anti-prediktif dönemlerde "abstain" kararı.'),

    ('Geirhos, R., et al. (2020). Shortcut learning in deep neural networks. Nature Machine Intelligence, 2(11), 665-673.',
     'https://doi.org/10.1038/s42256-020-00257-z',
     'DL modellerinin eğitim verisindeki kısayolları öğrendiğini gösterir. Projemizdeki modeller makro korelasyonları "kısayol" olarak öğrenip, test\'te bu kırılınca başarısız olmaktadır.'),

    ('Göçken, M., Özçalıcı, M., Boru, A., & Dosdoğru, A.T. (2016). Integrating metaheuristics and ANNs for stock price prediction. Expert Systems with Applications, 44, 320-331.',
     'https://doi.org/10.1016/j.eswa.2015.09.029',
     'Türkiye\'de BIST üzerine yapılan hibrit ANN çalışması. MC tuzağı veya anti-prediktif davranış ele alınmamıştır — projemizin bu boşluğu doldurduğunu gösterir.'),

    ('He, H., & Garcia, E.A. (2009). Learning from imbalanced data. IEEE Transactions on Knowledge and Data Engineering, 21(9), 1263-1284.',
     'https://doi.org/10.1109/TKDE.2008.239',
     'Sınıf dengesizliğinin kapsamlı incelemesi. class_weight=balanced çözümümüzün teorik temelini oluşturur.'),

    ('Kara, Y., Boyacıoğlu, M.A., & Baykan, Ö.K. (2011). Predicting direction of stock price index movement using ANNs and SVMs. Expert Systems with Applications, 38(5), 5311-5319.',
     'https://doi.org/10.1016/j.eswa.2010.10.027',
     'Türkiye\'de İMKB-100 üzerinde yapılan ilk kapsamlı ANN çalışması. BIST DL literatürünün başlangıç noktası olarak alıntılanmıştır.'),

    ('Krawczyk, B. (2016). Learning from imbalanced data: Open challenges and future directions. Progress in AI, 5(4), 221-232.',
     'https://doi.org/10.1007/s13748-016-0094-0',
     'Cost-sensitive learning yöntemlerinin dengesiz veride etkili olduğunu gösterir. MC tuzağı çözümümüzün metodolojik referansıdır.'),

    ('Sezer, O.B., Gudelek, M.U., & Ozbayoglu, A.M. (2020). Financial time series forecasting with deep learning: A systematic literature review. Applied Soft Computing, 90, 106181.',
     'https://doi.org/10.1016/j.asoc.2020.106181',
     '150+ DL çalışmasını kapsayan sistematik derleme. Finansal DL alanının genel haritasını çıkarır; anti-prediktif davranış hiçbir çalışmada raporlanmamıştır.'),

    ('Tashman, L.J. (2000). Out-of-sample tests of forecasting accuracy: An analysis and review. International Journal of Forecasting, 16(4), 437-450.',
     'https://doi.org/10.1016/S0169-2070(00)00065-0',
     'Walk-forward doğrulamanın standart yöntem olarak kabul edilmesini sağlayan temel makale. 7-fold WF v2 yaklaşımımızın doğrudan referansıdır.'),
]

for i, (ref, url, why) in enumerate(refs, 1):
    p = doc.add_paragraph()
    r = p.add_run('[{}] '.format(i))
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(ref)
    r2.font.size = Pt(10)
    
    # URL satiri
    p_url = doc.add_paragraph()
    r_url = p_url.add_run('    Erişim: {}'.format(url))
    r_url.font.size = Pt(9)
    r_url.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    
    # Neden alintilandi
    p_why = doc.add_paragraph()
    r_why_label = p_why.add_run('    Projemizle İlişkisi: ')
    r_why_label.bold = True
    r_why_label.font.size = Pt(9)
    r_why_label.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    r_why_text = p_why.add_run(why)
    r_why_text.font.size = Pt(9)
    r_why_text.italic = True
    r_why_text.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()  # bosluk

# ════════════════════════════════════════════
# KAYDET
# ════════════════════════════════════════════
out_path = os.path.join(OUT, 'MC_AWARE_Kisisel_Rapor_v2.docx')
doc.save(out_path)
sz = os.path.getsize(out_path) / 1024
print('Kaydedildi: {}'.format(out_path))
print('Boyut: {:.1f} KB'.format(sz))
print('Paragraf: {}'.format(len(doc.paragraphs)))
print('Tablo: {}'.format(len(doc.tables)))
print('Gorsel: {}'.format(len(doc.inline_shapes)))
