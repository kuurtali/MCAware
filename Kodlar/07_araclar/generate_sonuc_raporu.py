from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ========== SAYFA AYARLARI (SABLON: 2.5cm tum kenarlar) ==========
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5)   # FIX: 3cm -> 2.5cm
    s.right_margin = Cm(2.5)

# ========== FONT 12pt (SABLON: 12pt) ==========
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)  # FIX: 9pt -> 12pt
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(6)

for i in range(1, 4):
    hs = doc.styles['Heading ' + str(i)]
    hs.font.name = 'Arial'
    hs.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Siyah basliklar (resmi)
    hs.font.bold = True

# ========== GENEL BİLGİLER (FIX: eksik baslik eklendi) ==========
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TÜBİTAK 2209-A ÜNİVERSİTE ÖĞRENCİLERİ ARAŞTIRMA PROJELERİ\nDESTEKLEME PROGRAMI')
r.font.size = Pt(14)
r.font.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SONUÇ RAPORU')
r.font.size = Pt(14)
r.font.bold = True

doc.add_paragraph()

doc.add_heading('GENEL BİLGİLER', level=1)  # FIX: eksik baslik eklendi

# ========== TABLO 1: KESIN SABLON ALAN ADLARI ==========
t1 = doc.add_table(rows=4, cols=2, style='Table Grid')
t1.alignment = WD_TABLE_ALIGNMENT.CENTER

# FIX: Kesin sablon alan adlari kullaniliyor
data = [
    ('PROJENİN KONUSU', 'MC-AWARE: Derin Öğrenmede Anti-Prediktif Davranışın Çok Boyutlu Teşhisi'),
    ('PROJE YÜRÜTÜCÜSÜNÜN ADI', 'Mehmet Ali Kurt'),             # FIX: adi eklendi
    ('DANIŞMANIN ADI', 'Öğr. Gör. Elif Ayaz'),                  # FIX: adi eklendi
    ('PROJE BAŞLANGIÇ VE BİTİŞ TARİHLERİ', 'Başlangıç: Ekim 2025 — Bitiş: Haziran 2026')  # FIX: tam ad
]
for i, (k, v) in enumerate(data):
    t1.cell(i, 0).text = k
    t1.cell(i, 1).text = v
    for c in range(2):
        for r in t1.cell(i, c).paragraphs[0].runs:
            r.font.size = Pt(11)
            if c == 0:
                r.font.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Sonuç Raporu Formatı:')
r.font.bold = True
r.font.size = Pt(12)

doc.add_paragraph()

# ========== 1. GİRİŞ ==========
doc.add_heading('Giriş', level=1)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.0)
p.add_run(
    'Bu proje, TÜBİTAK 2209-A Lisans Araştırma Projeleri Destekleme Programı kapsamında, '
    'Borsa İstanbul (BIST) günlük yön tahmininde derin öğrenme modellerinin '
    'beklenmedik bir davranışını — anti-prediktif eğilimi — sistematik olarak tespit etmeyi '
    've çok boyutlu teşhisini gerçekleştirmeyi amaçlamaktadır.'
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.0)
p.add_run(
    'Anti-prediktif davranış, bir makine öğrenmesi modelinin rastgele tahminden (~%50) sistematik olarak '
    'daha kötü performans sergilemesi, yani tahminlerinin gerçek sonuçlarla negatif korelasyon göstermesi '
    'durumudur. Bu fenomen, modelin güçlü bir sinyal yakaladığını ancak bu sinyalin yönünün '
    'test döneminde tersine döndüğünü gösterir.'
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.0)
p.add_run(
    'Projede 6 farklı derin öğrenme mimarisi (BiLSTM, GRU, Conv1D, TCN, Transformer, SimpleRNN) ile '
    "700'den fazla konfigürasyon, 11 BIST hissesi, 27 varlık ve 130 CSV çıktı üretilmiştir. "
    'Türk akademik literatüründe, derin öğrenme modellerinin Çoğunluk Sınıfı (MC) tuzağını '
    'sistematik olarak belgeleyen ve anti-prediktif davranışı çok boyutlu teşhis eden ilk çalışmadır.'
)

# ========== 2. YAPILAN ÇALIŞMALAR ==========
doc.add_heading('Rapor Dönemlerinde Yapılan Çalışmalar', level=1)

# --- Faz 1 ---
p = doc.add_paragraph()
r = p.add_run('Faz 1: Prototipleme ve İlk Deneyler (Ekim–Aralık 2025)')
r.font.bold = True
r.font.size = Pt(12)

for item in [
    'BiLSTM v1-v3 prototiplerinin geliştirilmesi ve THYAO hissesi üzerinde test edilmesi.',
    'Makroekonomik değişkenlerin (USDTRY, Brent Petrol, TCMB Faiz) modele entegrasyonu.',
    'Çoğunluk Sınıfı (MC) tuzağının tespit edilmesi ve class_weight=balanced yöntemiyle çözülmesi.',
    'Anti-prediktif davranışın ilk kez gözlemlenmesi: model ~%40 doğruluk, flip ~%60 doğruluk.',
    'İlk CSV çıktılarının üretilmesi ve analiz pipeline oluşturulması.',
]:
    doc.add_paragraph(item, style='List Bullet')

# --- Faz 2 ---
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Faz 2: Ablasyon ve Diagnostik Deneyler (Ocak–Şubat 2026)')
r.font.bold = True
r.font.size = Pt(12)

for item in [
    'Özellik Grubu Ablasyonu: full_13 vs no_ext_10 karşılaştırması. Makro değişkenler çıkarıldığında anti-prediktif davranış tamamen ortadan kalkmıştır.',
    'Tekli Özellik Ablasyonu: Tek değişken çıkarıldığında etki gözlemlenmemiş, dış değişkenlerin sinerjik sahte korelasyon oluşturduğu kesinleşmiştir.',
    'IN_LEN Ablasyonu v1 ve v2: IN_LEN ≤ 5 anti-prediktif davranışı tetiklerken, IN_LEN = 10 tamamen ortadan kaldırmıştır.',
    'MI Kalibrasyon: Karşılıklı Bilgi (MI) skorlarının histogram yanlılığı içerdiği saptanmıştır.',
    'Etiket doğrulama ve veri bütünlüğü kontrolü gerçekleştirilmiştir.',
]:
    doc.add_paragraph(item, style='List Bullet')

# --- Faz 3 ---
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Faz 3: Çoklu Mimari Testi ve Walk-Forward Doğrulama (Mart–Nisan 2026)')
r.font.bold = True
r.font.size = Pt(12)

for item in [
    '6 farklı derin öğrenme mimarisinin (BiLSTM, GRU, Conv1D, TCN, Transformer, SimpleRNN) sistematik karşılaştırması yapılmıştır.',
    'McNemar testi ile mimariler arası istatistiksel fark analizi gerçekleştirilmiştir.',
    '7-fold Walk-Forward v1 ve v2 doğrulama uygulanmış, v2 ile validasyon verisi (val_data) düzeltmesi yapılmıştır.',
    'Dönemsel anti-prediktif harita çıkarılmıştır: Fold 2-3 (COVID-19 dönemi) = %100 anti-prediktif, Fold 1, 4, 5 (normal dönemler) = %0.',
    'Ensemble deneyleri (Hard Vote, Soft Vote) gerçekleştirilmiş, anti-prediktif davranışın ensemble seviyesinde de korunduğu gösterilmiştir.',
]:
    doc.add_paragraph(item, style='List Bullet')

# --- Faz 4 ---
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Faz 4: Çapraz Piyasa ve Çoklu Hisse Testleri (Nisan–Mayıs 2026)')
r.font.bold = True
r.font.size = Pt(12)

for item in [
    'NASDAQ (AAPL) vs BIST (THYAO) çapraz piyasa testi: BIST anti-prediktif davranış sergilerken, NASDAQ normal davranış göstermiştir.',
    '11 BIST hissesinde genişletilmiş test: 5 hisse anti-prediktif (THYAO, PGSUS, HEKTS, SASA, KRDMD), 6 hisse normal/nötr davranış.',
    'Sektörel ayrışma kesin olarak kanıtlanmıştır: havacılık ve spekülatif sektörler anti-prediktif, bankacılık ve otomotiv sektörleri normal.',
    'Bonferroni-düzeltilmiş p = 0.00012 ile istatistiksel anlamlılık doğrulanmıştır.',
    'BES fonları dahil sigorta ve holding sektörlü testler gerçekleştirilmiştir.',
]:
    doc.add_paragraph(item, style='List Bullet')

# --- Faz 5 ---
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Faz 5: Dokümantasyon ve Yayın Hazırlığı (Mayıs–Haziran 2026)')
r.font.bold = True
r.font.size = Pt(12)

for item in [
    'Kapsamlı literatür taraması: 25 bölüm, 30+ referans (Türk çalışmaları dahil).',
    'Araştırmanın kısıtları: 15 kısıt ve 15 gelecek çalışma iş paketi belirlenmiştir.',
    'Streamlit interaktif dashboard geliştirilmiştir (10 tab, TR/EN dil desteği).',
    'GitHub Pages statik dashboard oluşturulmuştur.',
    '37 rapor görseli ve 4 özel rapor görseli (G1-G4) hazırlanmıştır.',
    'Profesyonel dosya yapısı düzenlemesi ve GitHub deposuna aktarım (248 dosya) tamamlanmıştır.',
]:
    doc.add_paragraph(item, style='List Bullet')

# ========== 3. SONUÇ ==========
doc.add_heading('Sonuç', level=1)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.0)
p.add_run(
    'Proje kapsamında gerçekleştirilen kapsamlı deneysel çalışma, '
    'derin öğrenme modellerinin gelişmekte olan piyasalarda (BIST) sistematik olarak '
    'anti-prediktif davranış sergileyebildiğini kesin olarak kanıtlamıştır. '
    'Aşağıdaki tablo, projenin temel özgün katkılarını özetlemektedir:'
)

doc.add_paragraph()
t2 = doc.add_table(rows=1, cols=2, style='Table Grid')
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
t2.cell(0, 0).text = 'Bulgu'
t2.cell(0, 1).text = 'Kanıt'
for c in range(2):
    for r in t2.cell(0, c).paragraphs[0].runs:
        r.font.bold = True
        r.font.size = Pt(10)

findings = [
    ('Anti-prediktif davranış', '700+ konfigürasyon, 6 mimari, ~%40 model → ~%60 flip'),
    ('Sektörel ayrışma', 'Havacılık/spekülatif: %100, bankacılık: normal (p < 0.0001)'),
    ('Pencere uzunluğu etkisi', 'IN_LEN ≤ 5: tetikler, IN_LEN = 10: ortadan kaldırır'),
    ('Makro değişken mekanizması', 'USDTRY + Oil + TCMB çıkarılınca anomali tamamen yok'),
    ('Dönemsel harita', 'COVID (Fold 2-3): %100, Normal (Fold 1,4,5): %0'),
    ('Çapraz piyasa', 'BIST anti-prediktif, NASDAQ normal'),
    ('MC tuzağı çözümü', '700+ koşuda MC = 0 (hiçbir model çoğunluk sınıfına çökmedi)'),
    ('Falsifikasyon', '11 hisselik stres testi, Bonferroni p = 0.00012'),
    ('İstatistiksel anlamlılık', 'Bireysel p < 0.0001, Bonferroni-düzeltilmiş p = 0.00012'),
    ('Ensemble direnci', '6-mimari ensemble da anti-prediktif davranışı düzeltmedi'),
]
for f, k in findings:
    row = t2.add_row()
    row.cells[0].text = f
    row.cells[1].text = k
    for c in range(2):
        for r2 in row.cells[c].paragraphs[0].runs:
            r2.font.size = Pt(10)

doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.0)
p.add_run(
    'Sonuç olarak, projemiz derin öğrenme modellerinin finansal tahmin görevlerinde '
    '"her zaman daha iyi" olmadığını, aksine belirli makroekonomik koşullar altında '
    'sistematik olarak yanlış yöne tahmin yapabildiğini göstermiştir. '
    'Bu "anti-prediktif" davranış, literatürde henüz yeterince incelenmemiş '
    'kritik bir fenomendir ve projemiz bu boşluğu doldurmaktadır.'
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.0)
p.add_run(
    'Projenin gelecek çalışmalarında, 15 iş paketi olarak belirlenen '
    'rejim algılama, seçici tahmin (Selective Abstain), çapraz piyasa doğrulama '
    've duygu analizi entegrasyonu gibi konuların araştırılması hedeflenmektedir.'
)

# ========== 4. ÇIKTILAR ==========
doc.add_heading('Çıktılar (Yayınlar, Sunumlar vb.)', level=1)

outputs = [
    ('GitHub Deposu', 'https://github.com/kuurtali/Tubitak-2209A-MCAware — 248 dosya, 130 CSV çıktı'),
    ('Streamlit Dashboard', 'mcaware.streamlit.app — 10 tab, TR/EN, gerçek zamanlı interaktif araştırma paneli'),
    ('Akademik Doküman', 'Literatur_ve_Limitasyonlar.docx — 25 literatür bölümü + 15 kısıt + 15 gelecek çalışma iş paketi'),
    ('Rapor Görselleri', '37 görsel (PNG) + 4 özel rapor görseli (sektörel dağılım, fold haritası vb.)'),
    ('Veri Çıktıları', '130 CSV dosyası: 71 özet + 19 tahmin + 26 diagnostik + 14 eşik grid analizi'),
    ('R Scriptleri', '52 R scripti, 7 klasör altında organize (01_prototypes → 07_araclar)'),
]
for title, desc in outputs:
    p = doc.add_paragraph()
    r = p.add_run(title + ': ')
    r.font.bold = True
    p.add_run(desc)

# ========== 5. HARCAMALAR ==========
doc.add_heading('Proje ile İlgili Harcama Kalemleri Hakkında Ayrıntılı Bilgi', level=1)

p = doc.add_paragraph()
r = p.add_run('(Bu bölüm proje yürütücüsü tarafından doldurulacaktır.)')
r.italic = True

# ========== İMZA TABLOSU (SABLON FORMATI) ==========
doc.add_paragraph()
doc.add_paragraph()

t3 = doc.add_table(rows=3, cols=2, style='Table Grid')
t3.alignment = WD_TABLE_ALIGNMENT.CENTER

# FIX: Sablon formati - ust satir baslik
c00 = t3.cell(0, 0)
p00 = c00.paragraphs[0]
r00 = p00.add_run('PROJE YÜRÜTÜCÜSÜNÜN\nADI – SOYADI - İMZA')
r00.font.bold = True
r00.font.size = Pt(11)
p00.alignment = WD_ALIGN_PARAGRAPH.CENTER

c01 = t3.cell(0, 1)
p01 = c01.paragraphs[0]
r01 = p01.add_run('DANIŞMANIN\nADI – SOYADI - İMZA')
r01.font.bold = True
r01.font.size = Pt(11)
p01.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Isim satirlari
t3.cell(1, 0).text = 'Mehmet Ali Kurt'
t3.cell(1, 1).text = 'Öğr. Gör. Elif Ayaz'
for c in range(2):
    t3.cell(1, c).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t3.cell(1, c).paragraphs[0].runs:
        r.font.size = Pt(11)

# Bos imza satiri
t3.cell(2, 0).text = ''
t3.cell(2, 1).text = ''

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Tarih : Haziran 2026')

# ========== KAYDET ==========
outpath = 'Docs/TUBITAK_2209A_Sonuc_Raporu.docx'
doc.save(outpath)
sz = os.path.getsize(outpath) / 1024
print('Rapor: {:.1f} KB'.format(sz))
print('DUZELTMELER:')
print('  [x] Tablo 1 alan adlari = sablon ile ayni')
print('  [x] Font 12pt')
print('  [x] Margin 2.5cm (tum kenarlar)')
print('  [x] GENEL BILGILER basligi eklendi')
print('  [x] Imza tablosu sablon formatinda')
print('  [x] Bolum basliklari sablon diline uyumlastirild')
print('  [x] Harcama basligı sablondaki gibi uzun format')
