# -*- coding: utf-8 -*-
"""Update TUBITAK_2209A_Proje_Durumu.docx with 3 new graphics"""

import os
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTDIR = r"C:\Users\Kurt\Desktop\Proje\00_Tubitak\Gorseller"
DOCS = r"C:\Users\Kurt\Desktop\Proje\00_Tubitak\Docs"
docx_path = os.path.join(DOCS, "TUBITAK_2209A_Proje_Durumu.docx")

# Yedek al
now_str = datetime.now().strftime("%Y%m%d_%H%M%S")
backup_name = f"TUBITAK_2209A_Proje_Durumu_YEDEK_{now_str}.docx"
backup_path = os.path.join(DOCS, backup_name)
shutil.copy2(docx_path, backup_path)
print(f"[OK] Yedek: {backup_name}")

# Docx ac
doc = Document(docx_path)

# Remove duplicate styles to avoid KeyError
seen_ids = set()
styles_el = doc.styles.element
nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
for style_el in list(styles_el.findall(".//w:style", nsmap)):
    sid = style_el.get(qn("w:styleId"))
    if sid in seen_ids:
        styles_el.remove(style_el)
    else:
        seen_ids.add(sid)

def add_heading_safe(doc, text, level=1):
    p = doc.add_paragraph(text)
    try:
        p.style = f"Heading {level}"
    except KeyError:
        pPr = p._element.get_or_add_pPr()
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), f"Heading{level}")
        pPr.insert(0, pStyle)
    return p

# Sayfa sonu
doc.add_page_break()

# Ana baslik
heading = add_heading_safe(doc, "Ek Gorseller ve Analizler", level=1)
for run in heading.runs:
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    run.font.size = Pt(18)
    run.bold = True

# Tarih
date_para = doc.add_paragraph()
date_str = datetime.now().strftime("%d %B %Y, %H:%M")
dr = date_para.add_run(f"Ekleme Tarihi: {date_str}")
dr.font.size = Pt(10)
dr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
dr.italic = True
doc.add_paragraph("")

# === Gorsel 31: Proje Ozet Infografigi ===
h31 = add_heading_safe(doc, "Proje Ozet Infografigi", level=2)
for run in h31.runs:
    run.font.size = Pt(14)
    run.bold = True

p = doc.add_paragraph()
p.add_run(
    "Bu gorsel, MC-AWARE projesinin 26 deneyinin (22 ana + 4 faz-2) "
    "tum ana bulgularini tek bir sayfada ozetlemektedir. Her kart, "
    "bir deney kategorisini (MC Tuzagi Cozumu, Anti-Prediktif Davranis, "
    "Mimari Bagimsizlik, Walk-Forward Dogrulama, Klasik ML Negatif Kontrol, "
    "Feature Ablasyon, IN_LEN Kirilganlik, Sektorel Kontrast) temsil "
    "eder. Alt barda toplam istatistikler (26 deney, 6 DL mimari, "
    "378+ konfigurasyon, 30+ gorsel, 115+ CSV, 7 walk-forward fold) "
    "yer almaktadir."
).font.size = Pt(10)

doc.add_picture(
    os.path.join(OUTDIR, "31_Proje_Ozet_Infografik.png"), width=Inches(6.0)
)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

# === Gorsel 32: Concept Drift Zaman Serisi ===
h32 = add_heading_safe(doc, "Concept Drift Zaman Serisi", level=2)
for run in h32.runs:
    run.font.size = Pt(14)
    run.bold = True

p = doc.add_paragraph()
p.add_run(
    "Ust panel: USDTRY ve Petrol fiyatlarinin THYAO ile olan korelasyonlarinin "
    "egitim ve test setleri arasindaki kirilmasini (concept drift) gostermektedir. "
    "Kirmizi alan USDTRY, mavi alan Oil degiskeninin train-test korelasyon farkini "
    "(drift buyuklugunu) temsil eder. Ozellikle Fold 3-7 arasinda USDTRY "
    "korelasyonu dramatik olarak dusmekte, Oil korelasyonu ise isaret "
    "degistirmektedir. Alt panel: Her fold icin model dogrulugu (kirmizi), "
    "ters-yon dogrulugu (yesil) ve naive baseline (sari) karsilastirilmaktadir. "
    "Anti-prediktif foldlar (Fold 1, 2, 6) acik kirmizi arka plan ile "
    "isaretlenmistir. Bu grafik, korelasyon kirilmasinin anti-prediktif "
    "davranisla zamansal olarak ortusmesini dogrudan gostermektedir."
).font.size = Pt(10)

doc.add_picture(
    os.path.join(OUTDIR, "32_Concept_Drift_Timeline.png"), width=Inches(6.0)
)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

# === Gorsel 33: Metodoloji Pipeline ===
h33 = add_heading_safe(doc, "Metodoloji Pipeline Akis Semasi", level=2)
for run in h33.runs:
    run.font.size = Pt(14)
    run.bold = True

p = doc.add_paragraph()
p.add_run(
    "MC-AWARE projesinin uctan uca metodoloji akis semasi. "
    "Ust satir (Ana Hat): Ham veri toplama (Yahoo Finance, 2014-2023) -> "
    "Feature muhendisligi (13 ozellik: 10 teknik indikator + 3 makro degisken) -> "
    "MC-Aware on isleme (class_weight=balanced + MC-Aware loss ile lambda tuning) -> "
    "Derin ogrenme modelleri (6 mimari: BiLSTM, GRU, Conv1D, TCN, Transformer, RNN) -> "
    "Degerlendirme ve analiz (Walk-Forward CV, 7 fold, 5 seed, Flip-Acc metrigi). "
    "Alt satir (Dogrulama Katmani): Negatif kontrol (klasik ML baseline) -> "
    "Ablasyon testleri (feature, IN_LEN, seed) -> Istatistiksel dogrulama "
    "(McNemar, binomial CI) -> Sektorel genelleme (sigorta, holding, NASDAQ) -> "
    "Bulgular (anti-prediktif oruntusu, concept drift kaniti, EMH-uyumlu sonuc). "
    "Toplam 378+ konfigurasyon, 26 deney, 90/90 MC=0 basarisi."
).font.size = Pt(10)

doc.add_picture(
    os.path.join(OUTDIR, "33_Metodoloji_Pipeline.png"), width=Inches(6.0)
)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Kaydet
doc.save(docx_path)
print("[OK] TUBITAK_2209A_Proje_Durumu.docx guncellendi.")
print("TUM ISLEMLER TAMAMLANDI!")
